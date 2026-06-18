#!/usr/bin/env python3
"""Generate PROJECT_ROMAN_URDU.docx — project explanation in Roman Urdu."""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(PROJECT_ROOT, "PROJECT_ROMAN_URDU.docx")


def shade_cell(cell, hex_color="D9E2F3"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_color="4472C4"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, header_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    doc.add_paragraph("")
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc():
    doc = Document()

    title = doc.add_heading(
        "Prompt Injection aur Command Injection\n"
        "Big Data Platforms par Security Project",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Roman Urdu Manual — Project kya hai, hum ne kya kiya, har folder ki detail"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Team
    doc.add_heading("Team Members", level=1)
    add_table(doc, ["Roll Number", "Name"], [
        ("L1F22BSCS0122", "Nauman Irshad Ali Shah"),
        ("L1F22BSCS0929", "Saadia Lucman"),
        ("L1F22BSCS0939", "Syeda Aliza Wajid"),
        ("L1F22BSCS0909", "Hira Noor"),
    ])

    # 1. Project kya hai
    doc.add_heading("1. Yeh Project Kya Hai?", level=1)
    doc.add_paragraph(
        "Yeh hamara university security research project hai. Is mein hum ne do bari "
        "problems study ki hain jo aaj kal Big Data systems (jaise Hadoop aur Spark) aur "
        "AI chatbots (LLM apps) mein hoti hain."
    )
    add_bullets(doc, [
        "Command Injection — hacker server par shell commands chala sakta hai.",
        "Prompt Injection — hacker AI ko chhoti si hidden line se trick kar sakta hai.",
        "Hum ne real attack demos banaye, phir un par defense lagaya, aur akhir mein "
        "ek Security Chatbot banaya jo live check karta hai ke user ka message safe hai ya attack hai.",
        "Har query Apache Spark (PySpark) se log hoti hai taake hum analytics kar saken.",
    ])

    # 2. Hum ne kya kiya
    doc.add_heading("2. Hum Ne Kya Kiya?", level=1)
    add_bullets(doc, [
        "Hadoop aur Spark par command injection attacks demonstrate kiye (CVE numbers ke sath).",
        "RAG Poisoning, PIDP, P2SQL jaise prompt injection attacks banaye aur test kiye.",
        "Mirror Detector, CleanBase, aur UTDMF naam ka defense framework implement kiya.",
        "Flask server + HTML chatbot UI banaya — safe prompt par Small LLM jawab deta hai, "
        "attack par BLOCK karta hai aur reason dikhata hai.",
        "Research dashboard banaya jahan attacks, defenses aur project info hai.",
        "Poora project GitHub par publish kiya taake koi bhi download kar ke dekh sake.",
    ])

    # 3. Techniques
    doc.add_heading("3. Kaun Si Techniques Use Ki Hain?", level=1)
    add_table(doc, ["Area", "Technique", "Simple Matlab"], [
        ("Big Data", "Apache Hadoop", "Distributed file system — yahan CVE attack demo hai"),
        ("Big Data", "Apache Spark / PySpark", "Fast data processing — queries log hoti hain"),
        ("ML Defense", "Linear SVM + n-grams", "Mirror Detector — text mein injection dhundta hai"),
        ("ML Defense", "Semantic similarity", "CleanBase — poisoned RAG documents pakarta hai"),
        ("SQL Safety", "Whitelist + keyword block", "P2SQL Validator — DROP/DELETE rokta hai"),
        ("AI", "google/flan-t5-small", "Chhota LLM — sirf safe queries ka jawab"),
        ("Vector DB", "ChromaDB", "RAG poisoning attack demo ke liye"),
        ("Web", "Flask + HTML/CSS/JS", "Chatbot API aur user interface"),
        ("Language", "Python 3", "Saari attacks, defense aur chatbot code"),
    ], header_color="5B4B8A")

    # 4. Folder structure
    doc.add_heading("4. Har Folder Mein Kya Hai? (Detail)", level=1)
    doc.add_paragraph(
        "Neeche har folder aur file explain ki gayi hai — kya kaam karti hai aur "
        "project mein kya role hai. Run commands yahan nahi hain, sirf samajhne ke liye."
    )

    # attacks/
    doc.add_heading("Folder: attacks/", level=2)
    doc.add_paragraph(
        "Yeh main attack demos ka folder hai. Har script alag attack type dikhati hai. "
        "Educational purpose ke liye hai — real hack nahi, sirf concept demonstrate karta hai."
    )

    doc.add_heading("attacks/command_injection/", level=3)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "cve_2022_25168_demo.py",
            "Hadoop CVE-2022-25168 demo. Tar file ke andar dangerous file name se "
            "shell command inject hone ka concept dikhata hai. CVSS 9.8 — bohot dangerous bug tha.",
        ),
        (
            "spark_50239_demo.py",
            "Spark SPARK-50239 demo. spark-submit ki JavaOptions setting se command "
            "execution ka attack vector explain karta hai.",
        ),
    ])

    doc.add_heading("attacks/prompt_injection/", level=3)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "rag_poisoning.py",
            "RAG Poisoning attack. ChromaDB vector database mein aise documents daalta hai "
            "jin ke andar chhupi hui malicious instructions hon — jaise 'IGNORE ALL RULES'. "
            "Jab AI document search karta hai to poisoned text mil jata hai.",
        ),
        (
            "pidp_attack.py",
            "PIDP = Prompt Injection + Database Poisoning. Ek sath prompt trick aur "
            "database mein poisoned passages combine karta hai — double attack.",
        ),
        (
            "p2sql_attack.py",
            "P2SQL attack. User normal English mein kehta hai 'show records and drop table' — "
            "yeh SQL injection jaisa kaam karta hai LLM analytics interface par.",
        ),
    ])

    doc.add_heading("attacks/combined_attacks/", level=3)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "chain_attack.py",
            "Multiple attacks ek chain mein — pehle ek attack, phir doosra. "
            "Dikhata hai ke attacker step by step system tod sakta hai.",
        ),
        (
            "mitigation_test.py",
            "Attack ke baad defense test karta hai — dekhta hai ke hamari protections "
            "kaam kar rahi hain ya nahi.",
        ),
    ])

    # defense/
    doc.add_heading("Folder: defense/", level=2)
    doc.add_paragraph(
        "Yeh defense (protection) ka folder hai. Attacks ko detect aur block karne wala code yahan hai."
    )
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "mirror_detector.py",
            "Mirror Detector — machine learning (Linear SVM) se prompt injection text "
            "detect karta hai. Character n-grams use karta hai. Fast aur lightweight hai.",
        ),
        (
            "cleanbase_detector.py",
            "CleanBase — RAG knowledge base mein poisoned documents dhundta hai. "
            "Semantic clique analysis use karta hai — similar suspicious docs group karta hai.",
        ),
        (
            "utdmf_defense.py",
            "UTDMF = Unified Threat Detection framework. Mirror + CleanBase + P2SQL "
            "sab ko ek pipeline mein jodta hai. Main defense engine hai.",
        ),
        (
            "chatbot_engine.py",
            "Chatbot ka brain. User prompt aata hai → Mirror check → P2SQL check (agar data query ho) "
            "→ Keywords check → Spark log → SAFE ya BLOCK decision. Block par reasons batata hai.",
        ),
        (
            "llm_responder.py",
            "Small LLM handler. Sirf SAFE prompts ka jawab generate karta hai "
            "(google/flan-t5-small). Attack par LLM call nahi hoti.",
        ),
        (
            "__init__.py",
            "Python package file — defense folder ko module banata hai.",
        ),
    ])

    # backend/
    doc.add_heading("Folder: backend/", level=2)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "chat_server.py",
            "Flask web server. Chatbot API deta hai — /api/scan prompt check karta hai, "
            "/api/stats analytics deta hai, /api/project project info deta hai. "
            "Port 5000 par chalta hai.",
        ),
    ])

    # frontend/
    doc.add_heading("Folder: frontend/", level=2)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "chat.html",
            "Main chatbot UI. User yahan message likhta hai. Pipeline animation dikhti hai "
            "(Mirror → P2SQL → Spark). SAFE par green + LLM answer, BLOCK par red + reason.",
        ),
        (
            "index.html",
            "Research dashboard. Attacks, defenses, team info, commands table — "
            "presentation ke liye professional page.",
        ),
    ])

    # data/
    doc.add_heading("Folder: data/", level=2)
    doc.add_paragraph("Project ka sample data aur logs yahan store hote hain.")
    add_table(doc, ["Sub-folder / File", "Kya Hai"], [
        (
            "benign_documents/",
            "Normal, safe documents — fleet guidelines, vehicle maintenance info. "
            "RAG system ke liye clean data.",
        ),
        (
            "benign_documents/fleet_guidelines.txt",
            "Fleet management rules — safe reference document.",
        ),
        (
            "benign_documents/vehicle_maintenance.txt",
            "Vehicle maintenance schedule — safe reference document.",
        ),
        (
            "poisoned_documents/",
            "Attack demo ke liye malicious documents — andar hidden injection instructions.",
        ),
        (
            "poisoned_documents/poison_001.txt",
            "Pehla poisoned doc — chhupi hui bad instructions ke sath.",
        ),
        (
            "poisoned_documents/poison_002.txt",
            "Doosra poisoned doc — alag attack pattern.",
        ),
        (
            "query_logs/queries.jsonl",
            "Chatbot ki har query ka log — timestamp, safe/blocked, reasons. "
            "Spark analytics isi se kaam karti hai.",
        ),
        (
            "vector_db/",
            "ChromaDB database files. RAG poisoning demo isi par chalti hai — "
            "embeddings yahan save hoti hain.",
        ),
    ])

    # utils/
    doc.add_heading("Folder: utils/", level=2)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "spark_session.py",
            "PySpark session banata hai. Windows par bhi kaam kare is ke liye special settings. "
            "Java chahiye hoti hai.",
        ),
        (
            "spark_analytics.py",
            "Logged queries par Spark analytics — kitni safe, kitni blocked, statistics nikalta hai.",
        ),
        (
            "__init__.py",
            "Python package file.",
        ),
    ])

    # setup/
    doc.add_heading("Folder: setup/", level=2)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "hadoop_setup.sh",
            "Full Hadoop cluster install script (Linux/WSL). NameNode, DataNode, YARN setup steps.",
        ),
        (
            "spark_setup.sh",
            "Apache Spark install script. Big Data environment complete karne ke liye.",
        ),
    ])

    # scripts/
    doc.add_heading("Folder: scripts/", level=2)
    add_table(doc, ["File", "Kya Karti Hai"], [
        (
            "generate_manual_docx.py",
            "English Word manual banata hai — tables, flowchart, GitHub steps ke sath.",
        ),
        (
            "generate_roman_urdu_docx.py",
            "Yeh Roman Urdu document banata hai jo aap parh rahe hain.",
        ),
    ])

    # Root files
    doc.add_heading("Root Files (Project ki Main Files)", level=2)
    add_table(doc, ["File", "Kya Karti Hai"], [
        ("run_chatbot.py", "Security chatbot server start karta hai."),
        ("run_all.py", "Saari attack aur defense demos ek ek kar ke chalata hai."),
        ("verify_project.py", "Check karta hai sab kuch theek hai — PASS/FAIL report."),
        ("requirements.txt", "Python packages ki list — pyspark, flask, sklearn, chromadb, etc."),
        ("README.md", "GitHub par English documentation."),
        ("MANUAL.md", "Text format mein project manual."),
        ("MANUAL_copy.docx", "English Word manual — theory, flowchart, tables."),
        ("PROJECT_ROMAN_URDU.docx", "Yeh Roman Urdu document."),
        ("PROJECT_MANUAL.docx", "Full English project manual."),
        (".gitignore", "Git par kya upload nahi karna — cache, secrets, etc."),
    ])

    # Chatbot flow simple
    doc.add_heading("5. Chatbot Kaise Kaam Karta Hai? (Simple)", level=1)
    add_bullets(doc, [
        "User chatbot mein message likhta hai.",
        "Mirror Detector check karta hai — kya yeh prompt injection hai?",
        "Agar message database/SQL jaisa lage to P2SQL Validator check karta hai.",
        "Keyword scanner 'ignore', 'admin', 'drop table' jaise words dhundta hai.",
        "Har query Apache Spark mein log hoti hai.",
        "Agar sab theek → Small LLM (flan-t5-small) jawab deta hai.",
        "Agar attack → BLOCKED message + exact reason (kaun si layer ne roka).",
    ])

    # Attack vs defense summary
    doc.add_heading("6. Attacks vs Defense — Short Summary", level=1)
    add_table(doc, ["Attack", "Defense Jo Isko Rokti Hai"], [
        ("Hadoop Command Injection (CVE-2022-25168)", "Education + input validation concept"),
        ("Spark JavaOptions Injection (SPARK-50239)", "Config hardening concept"),
        ("RAG Poisoning", "CleanBase detector"),
        ("PIDP Attack", "CleanBase + Mirror Detector"),
        ("P2SQL Attack", "P2SQL Validator + Mirror Detector"),
        ("Combined Chain", "UTDMF full pipeline"),
        ("Koi bhi prompt injection", "Mirror Detector + Keywords + Chatbot block"),
    ], header_color="C55A11")

    doc.add_paragraph("")
    p = doc.add_paragraph("© 2026 — L1F22BSCS Project Team")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_doc()
